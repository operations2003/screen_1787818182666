import io
from docx import Document
from typing import Dict, Any
from app.services.text_cleaner import clean_extracted_text
from app.services.document_analyzer import analyze_document_quality

def parse_docx_bytes(docx_bytes: bytes) -> Dict[str, Any]:
    """
    Parses Word (.docx) documents using python-docx, extracting text from paragraphs and tables.
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
        chunks = []

        # Extract text from paragraphs
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                chunks.append(p.text.strip())

        # Extract text from tables without missing cell data
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_texts:
                    chunks.append(" | ".join(row_texts))

        extracted_text = "\n".join(chunks)
    except Exception as e:
        extracted_text = ""

    cleaned_text = clean_extracted_text(extracted_text)
    metrics = analyze_document_quality(cleaned_text, 1)

    return {
        "text": cleaned_text,
        "pageCount": 1,
        "extractionMethod": "python-docx",
        "ocrUsed": False,
        "textQuality": metrics["textQuality"],
        "characterCount": metrics["characterCount"],
        "wordCount": metrics["wordCount"]
    }
